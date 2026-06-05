"""
DOCX-Export für Businessplan.

Erweitert aus dem Verinaris-MVP — strukturierter, mit Tabellen statt
Bullet-Points wo angemessen.
"""

from io import BytesIO

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor

from app.services.businessplan.models import BusinessPlanResult


def export_docx(result: BusinessPlanResult) -> bytes:
    """Erzeugt einen DOCX-Businessplan als Bytes."""
    doc = Document()
    p = result.input

    # --- Titel + Header ---
    title = doc.add_heading(f"Businessplan – {p.startup_name}", level=0)
    sub = doc.add_paragraph()
    sub.add_run(
        f"Standort: {p.location or '—'} · Rechtsform: {p.legal_form or '—'} · "
        f"Gründer: {p.founder_name or '—'}"
    ).italic = True

    # --- 1. Executive Summary ---
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(result.summary)
    if result.llm_used and result.llm_used != "rule-based":
        ll = doc.add_paragraph()
        ll.add_run(f"Erzeugt mit: {result.llm_used}").italic = True

    # --- 2. Geschäftsidee ---
    doc.add_heading("2. Geschäftsidee", level=1)
    _add_field(doc, "Mission", p.mission)
    _add_field(doc, "Problem", p.problem)
    _add_field(doc, "Lösung", p.solution)
    _add_field(doc, "Alleinstellungsmerkmal (USP)", p.usp)
    _add_field(doc, "Zielkunden", p.target_customers)

    # --- 3. Geschäftsmodell ---
    doc.add_heading("3. Geschäftsmodell", level=1)
    _add_field(doc, "Umsatzmodell", p.revenue_model)
    _add_field(doc, "Vertriebskanäle", p.sales_channels)

    pricing_table = doc.add_table(rows=1, cols=2)
    pricing_table.style = "Light List Accent 1"
    pricing_table.rows[0].cells[0].text = "Tarif"
    pricing_table.rows[0].cells[1].text = "Preis / Monat"
    for label, value in [
        ("Basic", p.pricing_basic),
        ("Professional", p.pricing_pro),
        ("Enterprise", p.pricing_enterprise),
    ]:
        row = pricing_table.add_row().cells
        row[0].text = label
        row[1].text = f"{value:,.0f} €".replace(",", ".")

    # --- 4. Finanzplanung ---
    doc.add_heading("4. Finanzplanung (3-Jahres-Forecast)", level=1)
    fin_table = doc.add_table(rows=1, cols=6)
    fin_table.style = "Light Grid Accent 1"
    headers = ["Jahr", "Kunden", "Umsatz", "Fixkosten", "Marketing", "Ergebnis v.St."]
    for i, h in enumerate(headers):
        fin_table.rows[0].cells[i].text = h
        fin_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    for f in result.financials:
        row = fin_table.add_row().cells
        row[0].text = str(f.year)
        row[1].text = str(f.customers)
        row[2].text = _eur(f.total_revenue)
        row[3].text = _eur(f.fixed_costs)
        row[4].text = _eur(f.marketing)
        row[5].text = _eur(f.profit_before_tax)

    p_doc = doc.add_paragraph()
    p_doc.add_run(
        f"Kapitalbedarf: {_eur(p.required_capital)}  ·  "
        f"Eigenkapital: {_eur(p.founder_equity)}"
    ).bold = True

    # --- 5. Härtungs-Checks ---
    doc.add_heading("5. Härtung gegen IHK/HwK/Bank/BA", level=1)
    chk_table = doc.add_table(rows=1, cols=4)
    chk_table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Bereich", "Status", "Befund", "Empfehlung"]):
        chk_table.rows[0].cells[i].text = h
        chk_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    for c in result.checks:
        row = chk_table.add_row().cells
        row[0].text = c.area
        row[1].text = c.status
        row[2].text = c.finding
        row[3].text = c.recommendation

    # --- 6. Score-Card ---
    doc.add_heading("6. Businessplan-Score", level=1)
    score_table = doc.add_table(rows=1, cols=2)
    score_table.style = "Light List Accent 2"
    score_table.rows[0].cells[0].text = "Kriterium"
    score_table.rows[0].cells[1].text = "Score"
    for label, value in [
        ("Businessplan-Reifegrad", result.scores.business_plan_maturity),
        ("Bankenfähigkeit", result.scores.bankability),
        ("Förderfähigkeit", result.scores.fundability),
        ("Investorenfähigkeit", result.scores.investability),
    ]:
        row = score_table.add_row().cells
        row[0].text = label
        row[1].text = f"{value}/100"

    # --- 7. Fördermittel ---
    doc.add_heading("7. Fördermittel-Kandidaten", level=1)
    for fm in result.funding:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(f"{fm.name} ({fm.fit})")
        run.bold = True
        para.add_run(f": {fm.why}")
        para.add_run(f"  → Nächster Schritt: {fm.next_step}").italic = True

    # --- 8. Risiken ---
    if p.risks:
        doc.add_heading("8. Risiken", level=1)
        doc.add_paragraph(p.risks)

    # --- 9. Nächste Maßnahmen ---
    doc.add_heading("9. Empfohlene nächste Maßnahmen", level=1)
    next_steps = [
        "Pilotkunden und belastbare Referenzfälle definieren.",
        "Compliance-Paket mit AVV, TOMs, Löschkonzept und KI-Risikoklassifizierung erstellen.",
        "Vertriebsannahmen mit konkretem Funnel und Conversion-Raten hinterlegen.",
        "Liquiditätsplanung monatlich für 24 Monate ergänzen.",
        "Fördermittel vor Antragstellung tagesaktuell prüfen.",
    ]
    for s in next_steps:
        doc.add_paragraph(s, style="List Bullet")

    # --- Disclaimer ---
    doc.add_paragraph()
    disc = doc.add_paragraph()
    disc.add_run(
        "Hinweis: Dieses Dokument ersetzt keine Steuer-, Rechts- oder "
        "Fördermittelberatung. Es dient der strukturierten Vorbereitung "
        "für IHK/HwK, Banken, BA und Fördermittelgespräche."
    ).italic = True

    # Bytes zurückgeben
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# --- Helpers ---

def _add_field(doc: Document, label: str, text: str) -> None:
    if not text:
        return
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(text)


def _eur(value: float) -> str:
    return f"{value:,.0f} €".replace(",", ".")
